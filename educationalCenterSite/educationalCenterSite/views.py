from django.http import HttpResponseRedirect
from django.shortcuts import render
from django.http import HttpResponse
from django.utils.encoding import escape_uri_path
from django.conf import settings
from .forms import *

def index(request):
    return HttpResponseRedirect('/login')

def information(request):
    if not request.COOKIES.get('id'):
        return HttpResponseRedirect('/login')
    else:
        import sqlite3
        con = sqlite3.connect('educationalDate.db')
        cur = con.cursor()
        isAdmin = cur.execute(f"SELECT admin FROM users WHERE id={request.COOKIES.get('id')}").fetchone()[0]
        if str(isAdmin) == '1':
            return render(request, 'information.html')
        else:
            return render(request, 'information1.html')

def users(request):
    if not request.COOKIES.get('id'):
        return HttpResponseRedirect('/login')
    else:
        import sqlite3
        con = sqlite3.connect('educationalDate.db')
        cur = con.cursor()
        isAdmin = cur.execute(f"SELECT admin FROM users WHERE id={request.COOKIES.get('id')}").fetchone()[0]
        if str(isAdmin) != '1':
            return HttpResponseRedirect('/information')
    if request.method == 'POST':
        allUsers = cur.execute('SELECT id FROM users').fetchall()
        listData = [int(i[0])-1 for i in list(request.POST.keys())[1:]]
        for data in allUsers:
            if data[0] in listData:
                cur.execute(f'''UPDATE users SET admin = 1 WHERE id = {data[0]}''')
            else:
                cur.execute(f'''UPDATE users SET admin = 0 WHERE id = {data[0]}''')
        con.commit()
        return HttpResponseRedirect('/information/users')
    else:
        allUsers = cur.execute('SELECT * FROM users').fetchall()
        allUser = []
        for user in range(len(allUsers)):
            users1 = list(allUsers[user])
            users1[3] = UsersAdmin(initial={'admin':True}, prefix=user+1) if users1[3] == 1 \
                else UsersAdmin(prefix=user+1)
            allUser.append(users1)
        return render(request, 'users.html', {'allUsers': allUser})

def login(request):
    out = ''
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            import sqlite3
            con = sqlite3.connect('educationalDate.db')
            cur = con.cursor()
            name = form.cleaned_data['name']
            password = form.cleaned_data['password']
            passwordTrue = cur.execute(f'''SELECT id, password from users WHERE name = "{name}" ''').fetchall()
            if passwordTrue == []:
                out = 'Неверный логин'
            else:
                id = passwordTrue[0][0]
                if passwordTrue[0][1] == password:
                    outt = HttpResponseRedirect('/information')
                    outt.set_cookie("id", id, max_age=7*24*60*60)
                    return outt
                else:
                    out = 'Неверный пароль'
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form, 'out': out})

def employees(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    allEmployees = cur.execute('SELECT * FROM employees').fetchall()
    return render(request, 'employees.html', {'allEmployees': allEmployees})

def groupReports(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    allGroupReports = cur.execute('SELECT * FROM groupReports').fetchall()
    return render(request, 'groupReports.html', {'allGroupReports': allGroupReports})

def studyingProgramTranslate(listDate, x = 0):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    if x == 0:
        boolFields = ['Нет', 'Да']
        resulChoose = ['Зачет', 'Незачет', 'Пропуск']
        listDate[3] = boolFields[int(listDate[3])]
        listDate[8] = boolFields[int(listDate[8])]
        listDate[9] = boolFields[int(listDate[9])]
        listDate[10] = boolFields[int(listDate[10])]
        listDate[11] = boolFields[int(listDate[11])]
        listDate[12] = boolFields[int(listDate[12])]
        listDate[13] = boolFields[int(listDate[13])]
        listDate[14] = resulChoose[listDate[14]]
        return listDate
    else:
        nameGroup = cur.execute(f'select nameGroup from groupReports where id = {listDate[4]}').fetchone()[0]
        listDate.insert(5, nameGroup)
        return listDate

def studyingProgram(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    allStudyingProgram = cur.execute('SELECT * FROM studyingPrograms').fetchall()
    for i in range(len(allStudyingProgram)):
        allStudyingProgram[i] = studyingProgramTranslate(list(allStudyingProgram[i]))
    return render(request, 'studyingProgram.html', {'allStudyingProgram': allStudyingProgram})

def applicationsTranslate(listData, idUser = -1):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    if idUser == -1:
        boolVal = ['Нет', 'Да']
        stat = ['Новый', 'Выставлен счет', 'Обучение частично', 'Выполнена']
        listUser = [i[0] for i in cur.execute('SELECT name FROM users').fetchall()]
        listData[4] = stat[int(listData[4])]
        listData[5] = boolVal[int(listData[5])]
        listData[6] = boolVal[int(listData[6])]
        listData[8] = listUser[int(listData[8])]
        return listData
    elif idUser == -2:
        stat = ['Новый', 'Выставлен счет', 'Обучение частично', 'Выполнена']
        listUser = [i[0] for i in cur.execute('SELECT name FROM users').fetchall()]
        listData[2] = stat[int(listData[2])]
        listData[4] = listUser[int(listData[4])]
        return listData
    else:
        import datetime
        listData.insert(3, datetime.datetime.now().strftime('%d.%m.%Y %X'))
        listData.insert(4, 0)
        listData.append(0)
        listData.append(idUser)
        listData.append(listData[0])
        return listData

def protocolsTranslate(listData, idUser = -1):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    import datetime
    nameProgram = cur.execute(f'SELECT name FROM studyingPrograms where id = {listData[1]}').fetchone()[0]
    listData.insert(2, nameProgram)
    listData.insert(3, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'))
    listData.insert(4, 1)
    listData.append(idUser)
    return listData

def clients(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    allClients = cur.execute('SELECT * FROM clients').fetchall()
    return render(request, 'clients.html', {'allClients': allClients})
    
def add(request):
    listForm = [EmployeesForm, GroupReportsForm, ClientsForm, StudyingProgramForm, None, None]
    idForm = int(request.GET.get('id-form'))
    if request.method == 'POST':
        if not listForm[idForm](request.POST).is_valid():
            return render(request, 'add.html', {'form': listForm[idForm](), 'invalid': 'Заполните все поля'})
        import sqlite3
        con = sqlite3.connect('educationalDate.db')
        cur = con.cursor()
        addRequest = [
            ['employees', '?,?,?'],
            ['groupReports', '?,?'],
            ['clients', '?,?,?,?,?'],
            ['studyingPrograms', '?,?,?,?,?,?,?,?,?,?,?,?,?,?,?'],
            ['applications', '?,?,?,?,?,?,?,?,?,?'],
            None
        ]
        id = len(cur.execute(f"SELECT * FROM {addRequest[idForm][0]}").fetchall())
        if idForm == 3:
            listDate = tuple(studyingProgramTranslate(([id] + list(request.POST.values())[1:]), 1))
        else:
            listDate = tuple(([id] + list(request.POST.values())[1:]))
        cur.execute(f'''INSERT INTO {addRequest[idForm][0]} VALUES({addRequest[idForm][1]})''', listDate)
        con.commit()
        red = ['/information/employees', '/information/group-reports',
               '/applications/clients', '/information/studying-program',
               f'/applications/list?id={id}', f'/protocols/clients?id={id}']
        return HttpResponseRedirect(red[idForm])
    else:
        form = listForm[idForm]()
    return render(request, 'add.html', {'form': form})

def application(request):
    if not request.COOKIES.get('id'):
        return HttpResponseRedirect('/login')
    return render(request, 'application.html')

def edit(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    listForm = [EmployeesForm, GroupReportsForm, ClientsForm, StudyingProgramForm]
    idForm = int(request.GET.get('id-form'))
    idRow = int(request.GET.get('id-row'))
    addRequest = ['employees', 'groupReports', 'clients', 'studyingPrograms']
    if request.method == 'POST':
        listValue = list(request.POST.values())[1:]
        if not listForm[idForm](request.POST).is_valid():
            return render(request, 'edit.html', {'form': listForm[idForm](my_arg=listValue), 'invalid': 'Заполните все поля'})
        if idForm == 3:
            nameGroup = cur.execute(f'select nameGroup from groupReports where id = {listValue[3]}').fetchone()[0]
            listValue[3] = nameGroup
        listKey = list(request.POST.keys())[1:]
        outReq = f'UPDATE {addRequest[idForm]} SET '
        for dataOne in listKey:
            outReq += f'{dataOne} = ?, '
        outReq = outReq[:-2] + f' WHERE id = {idRow}'
        cur.execute(outReq, listValue)
        con.commit()
        red = ['/information/employees', '/information/group-reports', '/applications/clients',
               '/information/studying-program']
        return HttpResponseRedirect(red[idForm])
    else:
        listValue = cur.execute(f'SELECT * from {addRequest[idForm]} WHERE id = {idRow}').fetchone()
        listValue = list(listValue)
        listValue = listValue[1:]
        if idForm == 3:
            listValue = listValue[:4]+listValue[5:]
        form = listForm[idForm](my_arg=listValue)
    return render(request, 'edit.html', {'form': form})

def applicationsList(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    allApplications = cur.execute('SELECT * FROM applications').fetchall()
    nList = []
    for i in range(len(allApplications)):
        listApl = applicationsTranslate(list(allApplications[i]))
        listApl[1] = listApl[0] + 1
        nList.append(listApl)
    return render(request, 'applicationsList.html', {'allApplications': nList})

def listClientsNoAdmin(request):
    z = request.GET.get('idApplication')
    idApplication = -1 if z == "new" else int(z)
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    
    boolFields = ['Нет', 'Да']
    
    def retRender(date):
        form1 = ApplicationsFormNoAdmin(my_arg=date[0])
        
        nDate = date[3:]
        formsClients1 = []
        for i in range(len(nDate) // 6):
            listValues = nDate[i * 6:(i + 1) * 6]
            formsClients1.append(ClientListForm(my_arg=listValues, prefix=i))
        
        allApplications2 = list(cur.execute(f'SELECT * FROM applications where id = {idApplication}').fetchone())
        allApplications3 = applicationsTranslate([allApplications2[0]] +
                                                 allApplications2[3:5] + allApplications2[7:], -2)
        allApplications3 = allApplications3 + [boolFields[int(allApplications2[5])],
                                               boolFields[int(allApplications2[6])]]
        return render(request, 'listClients.html',
                      {'form': form1, 'formsClients': formsClients1, 'allApplications': allApplications3})
    
    def retNewRender(date):
        import datetime
        form1 = ApplicationsFormNoAdmin(my_arg=date[0])
        
        nDate = date[3:]
        formsClients1 = []
        for i in range(len(nDate) // 6):
            listValues = nDate[i * 6:(i + 1) * 6]
            formsClients1.append(ClientListForm(my_arg=listValues, prefix=i))
        
        id = len(cur.execute(f"SELECT * FROM applications").fetchall())
        allApplications3 = applicationsTranslate(
            [id, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'),
             0, 0, request.COOKIES.get('id')], -2)
        allApplications3 = allApplications3 + [boolFields[0],
                                               boolFields[0]]
        return render(request, 'listClients.html',
                      {'form': form1, 'formsClients': formsClients1, 'allApplications': allApplications3})
    
    if request.method == 'POST':
        listData = list(request.POST.values())[1:]
        applications = listData[0]
        value = listData[-1]
        listData = listData[:-1]
        
        if value == "0":
            listData = listData[1:]
            nameClient = cur.execute(f'SELECT name FROM clients where id = {applications[0]}').fetchone()[0]
            applications.insert(1, nameClient)
            
            if idApplication == -1:
                id = len(cur.execute(f"SELECT * FROM applications").fetchall())
                listData2 = tuple(
                    applicationsTranslate(([id] + applications + [0, 0]), request.COOKIES.get('id')))
                cur.execute(f'''INSERT INTO applications VALUES(?,?,?,?,?,?,?,?,?,?)''', listData2)
                
                countTrainee = 0
                for i in range(len(listData) // 6):
                    listValue = listData[i * 6:(i + 1) * 6]
                    idTrainee = i
                    
                    cur.execute(f'INSERT INTO trainee values '
                                + f"({id}, '{listValue[0]}', '{listValue[1]}', "
                                  f"'{listValue[2]}', '{listValue[3]}', '{listValue[4]}', "
                                  f"'{listValue[5]}', {idTrainee})")
                    countTrainee += 1
                idApplication = id
            else:
                cur.execute(
                    f'update applications set idClient = ?, nameClient = ? where id = {idApplication}',
                    tuple(applications))
                
                countTrainee = 0
                cur.execute(f"DELETE FROM trainee WHERE id={idApplication}")
                for i in range(len(listData) // 6):
                    listValue = listData[i * 6:(i + 1) * 6]
                    idr = cur.execute(f'SELECT MAX(idInd) FROM trainee WHERE id = {idApplication}').fetchone()[0]
                    idTrainee = int(idr if isinstance(idr, int) else -1) + 1
                    cur.execute(f'INSERT INTO trainee values '
                                + f"({idApplication}, '{listValue[0]}', '{listValue[1]}', "
                                + f"'{listValue[2]}', '{listValue[3]}', '{listValue[4]}', '{listValue[5]}', {idTrainee})")
                    countTrainee += 1
                cur.execute(f"UPDATE applications SET countEmployees = '{countTrainee}' WHERE id = {idApplication}")
            con.commit()
            return HttpResponseRedirect(f'/applications/list/clients?idApplication={idApplication}')
        elif value == "1":
            for _ in range(6):
                listData.append(None)
        elif int(value) > 1:
            idDel = int(value) - 3
            listData = listData[:6 * idDel + 3] + listData[6 * (idDel + 1) + 3:]
        else:
            idDel = abs(int(value)) - 1
            listData += listData[6 * idDel + 3:6 * (idDel + 1) + 3]
        if idApplication == -1:
            return retNewRender(listData)
        return retRender(listData)
    else:
        if idApplication == -1:
            import datetime
            id = len(cur.execute(f"SELECT * FROM applications").fetchall())
            form = ApplicationsFormNoAdmin()
            applicationsList2 = applicationsTranslate(
                [id, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'),
                 0, 0, request.COOKIES.get('id')], -2)
            
            return render(request, 'listClients.html',
                          {'form': form, 'formsClients': [],
                           'allApplications': applicationsList2})
        allApplications1 = list(cur.execute(f'SELECT * FROM applications where id = {idApplication}').fetchone())
        
        allApplicationsForm = [allApplications1[1]]
        allApplications = applicationsTranslate([allApplications1[0]] + allApplications1[3:5] + allApplications1[7:], -2)
        allApplications = allApplications + [boolFields[int(allApplications1[5])], boolFields[int(allApplications1[6])]]
        form = ApplicationsFormNoAdmin(my_arg=allApplicationsForm)
        
        listClient = list(cur.execute(f'SELECT * FROM trainee where id = {allApplications1[-1]}').fetchall())
        formsClients = []
        
        for i in range(len(listClient)):
            formsClients.append(ClientListForm(my_arg=listClient[i][1:], prefix=i))
        return render(request, 'listClients1.html', {'form': form, 'formsClients': formsClients,
                                                    'allApplications': allApplications})

def listClients(request):
    z = request.GET.get('idApplication')
    idApplication = -1 if z == "new" else int(z)
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    
    isAdmin = cur.execute(f"SELECT admin FROM users WHERE id={request.COOKIES.get('id')}").fetchone()[0]
    if str(isAdmin) != '1':
        return listClientsNoAdmin(request)
    def retRender(date):
        form1 = ApplicationsForm(my_arg=date[:3])
        
        nDate = date[3:]
        formsClients1 = []
        for i in range(len(nDate) // 6):
            listValues = nDate[i * 6:(i + 1) * 6]
            formsClients1.append(ClientListForm(my_arg=listValues, prefix=i))
        
        allApplications1 = list(cur.execute(f'SELECT * FROM applications where id = {idApplication}').fetchone())
        allApplications1 = applicationsTranslate([allApplications1[0]] +
                                                 allApplications1[3:5] + allApplications1[7:], -2)
        
        return render(request, 'listClients.html',
                      {'form': form1, 'formsClients': formsClients1, 'allApplications': allApplications1})
    def retNewRender(date):
        import datetime
        form1 = ApplicationsForm(my_arg=date[:3])
        
        nDate = date[3:]
        formsClients1 = []
        for i in range(len(nDate) // 6):
            listValues = nDate[i * 6:(i + 1) * 6]
            formsClients1.append(ClientListForm(my_arg=listValues, prefix=i))
        
        id = len(cur.execute(f"SELECT * FROM applications").fetchall())
        allApplications1 = applicationsTranslate([id, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'),
                                                  0, 0, request.COOKIES.get('id')], -2)
        
        return render(request, 'listClients.html',
                      {'form': form1, 'formsClients': formsClients1, 'allApplications': allApplications1})
    
    if request.method == 'POST':
        listData = list(request.POST.values())[1:]
        applications = listData[:3]
        value = listData[-1]
        listData = listData[:-1]
    
        if value == "0":
            listData = listData[3:]
            nameClient = cur.execute(f'SELECT name FROM clients where id = {applications[0]}').fetchone()[0]
            applications.insert(1, nameClient)
            
            if idApplication == -1:
                id = len(cur.execute(f"SELECT * FROM applications").fetchall())
                listData2 = tuple(
                    applicationsTranslate(([id] + applications), request.COOKIES.get('id')))
                cur.execute(f'''INSERT INTO applications VALUES(?,?,?,?,?,?,?,?,?,?)''', listData2)
                
                countTrainee = 0
                for i in range(len(listData) // 6):
                    listValue = listData[i * 6:(i + 1) * 6]
                    idTrainee = i
                    
                    cur.execute(f'INSERT INTO trainee values '
                                + f"({id}, '{listValue[0]}', '{listValue[1]}', "
                                f"'{listValue[2]}', '{listValue[3]}', '{listValue[4]}', "
                                f"'{listValue[5]}', {idTrainee})")
                    countTrainee += 1
                idApplication = id
            else:
                cur.execute(
                    f'update applications set idClient = ?, nameClient = ?, postpaid = ?, receiptPayment = ? where id = {idApplication}',
                    tuple(applications))
                
                countTrainee = 0
                cur.execute(f"DELETE FROM trainee WHERE id={idApplication}")
                for i in range(len(listData) // 6):
                    listValue = listData[i * 6:(i + 1) * 6]
                    idr = cur.execute(f'SELECT MAX(idInd) FROM trainee WHERE id = {idApplication}').fetchone()[0]
                    idTrainee = int(idr if isinstance(idr, int) else -1) + 1
                    cur.execute(f'INSERT INTO trainee values '
                     +f"({idApplication}, '{listValue[0]}', '{listValue[1]}', "
                     +f"'{listValue[2]}', '{listValue[3]}', '{listValue[4]}', '{listValue[5]}', {idTrainee})")
                    countTrainee += 1
                cur.execute(f"UPDATE applications SET countEmployees = '{countTrainee}' WHERE id = {idApplication}")
            con.commit()
            return HttpResponseRedirect(f'/applications/list/clients?idApplication={idApplication}')
        elif value == "1":
            for _ in range(6):
                listData.append(None)
        elif int(value) > 1:
            idDel = int(value) - 3
            listData = listData[:6*idDel+3] + listData[6*(idDel+1)+3:]
        else:
            idDel = abs(int(value)) - 1
            listData += listData[6 * idDel + 3:6 * (idDel + 1) + 3]
        if idApplication == -1:
            return retNewRender(listData)
        return retRender(listData)
    else:
        if idApplication == -1:
            import datetime
            id = len(cur.execute(f"SELECT * FROM applications").fetchall())
            form = ApplicationsForm()
            applicationsList2 = applicationsTranslate([id, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'),
                                 0, 0, request.COOKIES.get('id')], -2)
            
            return render(request, 'listClients.html',
                          {'form': form, 'formsClients': [],
                           'allApplications': applicationsList2})
        allApplications = list(cur.execute(f'SELECT * FROM applications where id = {idApplication}').fetchone())
        
        allApplicationsForm = [allApplications[1]]+allApplications[5:7]
        allApplications = applicationsTranslate([allApplications[0]]+allApplications[3:5]+allApplications[7:], -2)
        
        form = ApplicationsForm(my_arg=allApplicationsForm)
        
        listClient = list(cur.execute(f'SELECT * FROM trainee where id = {allApplications[-1]}').fetchall())
        formsClients = []
        
        for i in range(len(listClient)):
            formsClients.append(ClientListForm(my_arg=listClient[i][1:], prefix=i))
        return render(request, 'listClients.html', {'form': form, 'formsClients': formsClients,
                                                    'allApplications': allApplications})
    
def installAccount(request, id):
    import docx, sqlite3, json, datetime
    doc = docx.Document()
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    
    stat = cur.execute(f"select status from applications where id = {id}").fetchone()[0]
    if str(stat) == '0':
        cur.execute(f'''update applications set status = 1 where id = {id}''')
    
    applicationOne = cur.execute(f'SELECT * FROM applications where id = {id}').fetchone()
    client = cur.execute(f'SELECT * FROM clients where id = {applicationOne[-1]}').fetchone()
    par1 = doc.add_paragraph(f'Заказчик: {client[1]}').bold = True
    par2 = doc.add_paragraph(f'Юридический адрес ЗАКАЗЧИКА: {client[-1]}')
    par3 = doc.add_paragraph(f'ИНН ')
    s1 = par3.add_run(str(client[2]))
    s1.bold = True
    par3.add_run(f' КПП ')
    s2 = par3.add_run(str(client[3]))
    s2.bold = True
    par4 = doc.add_paragraph(f'Исполнитель: ')
    par4.bold = True
    with open('company.json') as json_file:
        data = json.load(json_file)
        par4.add_run(data['name'])
        par5 = doc.add_paragraph(f"Адрес {data['address']}")
        par6 = doc.add_paragraph(f"ИНН/КПП {data['inn']} / {data['kpp']}")
        par7 = doc.add_paragraph(f"Бик {data['bankAccount']}")
        supervisor = data['supervisor']
    allUsers = cur.execute(f'SELECT prefStudyingProgram FROM trainee WHERE id={id}').fetchall()
    allPrograms = [i[0] for i in allUsers]
    reqText = 'SELECT name, price, prefix, id FROM studyingPrograms WHERE '
    for i in allPrograms:
        reqText += f"id = '{i}' or "
    reqText = reqText[:-4]
    listPrograms = cur.execute(reqText).fetchall()
    setPrograms = set(allPrograms)
    table = doc.add_table(rows=len(setPrograms) + 2, cols=6)
    table.style = 'Table Grid'
    summ = 0
    for program in range(len(listPrograms)):
        table.cell(program+1, 0).text = str(program+1)
        table.cell(program+1, 1).text = listPrograms[program][0]
        table.cell(program+1, 2).text = 'чел.'
        countProg = allPrograms.count(str(listPrograms[program][-1]))
        table.cell(program+1, 3).text = str(countProg)
        table.cell(program+1, 4).text = str(listPrograms[program][1])
        table.cell(program+1, 5).text = str(int(listPrograms[program][1])*countProg)
        summ += int(listPrograms[program][1])*countProg
    table.cell(len(setPrograms)+1, 0).merge(table.cell(len(setPrograms)+1, 4))
    table.cell(len(setPrograms)+1, 0).text = 'Итого'
    table.cell(len(setPrograms)+1, 5).text = str(summ)
    table.cell(0, 0).text = ''
    table.cell(0, 1).text = 'Наименование товара'
    table.cell(0, 2).text = 'Единица измерения'
    table.cell(0, 3).text = 'Количество'
    table.cell(0, 4).text = 'Цена, руб.'
    table.cell(0, 5).text = 'Сумма, руб.'
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = ("attachment; filename="
                                       + escape_uri_path(f'счёт {applicationOne[2]} от {datetime.datetime.now(settings.TIME_DELTA).strftime("%d.%m.%Y %X")}.docx'))
    doc.save(response)
    con.commit()
    return response

def installProtocols(request, id):
    import docx, sqlite3
    doc = docx.Document()
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    
    def transProt(word, idWord):
        if idWord == 0:
            return cur.execute(f'select name from clients where id = {word}').fetchone()[0]
        stat = ['Зачет', 'Незачет', 'Пропуск']
        return stat[word]
    
    protocol = cur.execute(f'SELECT * FROM protocols WHERE id={id}').fetchone()
    program = cur.execute(f'SELECT * FROM studyingPrograms WHERE id={protocol[1]}').fetchone()
    protocolClients = cur.execute(f'SELECT * FROM protocolsClients WHERE idProtocols={id}').fetchall()
    
    name = f'Протокол №{id} от {protocol[3]}'
    h = doc.add_paragraph(name)
    h.bold = True
    h.alignment = 0
    doc.add_paragraph(program[6])
    doc.add_paragraph(protocol[6])
    #table
    listN = ['ФИО', 'Клиент', 'СНИЛС', 'Номер реестра',
             'Результат проведения', 'Подпись']
    listCol = []
    col1 = 1
    for prog in range(len(program[8:])):
        if str(program[8+prog]) == '1':
            col1 += 1
            listCol.append(listN[prog])
    table = doc.add_table(rows=len(protocolClients)+1, cols=col1)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = ''
    for col in range(len(listCol)):
        table.cell(0, col+1).text = listCol[col]
    for client in range(len(protocolClients)):
        for col in range(len(listCol)+1):
            if col == len(listCol):
                table.cell(client + 1, 0).text = protocol[7] + f'{(protocolClients[client][0]+1):02d}'
            elif listCol[col] == 'ФИО':
                table.cell(client + 1, col+1).text = str(protocolClients[client][2])
            elif listCol[col] == 'Клиент':
                table.cell(client + 1, col+1).text = transProt(protocolClients[client][3],
                                                               0)
            elif listCol[col] == 'СНИЛС':
                table.cell(client + 1, col+1).text = str(protocolClients[client][4])
            elif listCol[col] == 'Номер реестра':
                table.cell(client + 1, col+1).text = ''
            elif listCol[col] == 'Результат проведения':
                table.cell(client + 1, col+1).text = transProt(protocolClients[client][5],
                                                               1)
            elif listCol[col] == 'Подпись':
                table.cell(client + 1, col+1).text = ''
    response = HttpResponse(content_type='application/vnd.openxmlformats-'
                                         'officedocument.wordprocessingml.document')
    response['Content-Disposition'] = ("attachment; filename="
                                       + escape_uri_path(
                f'{name}.docx'))
    doc.save(response)
    con.commit()
    return response

def installProtocolsXML(request, id):
    import sqlite3, json
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    bodyXml = '''<?xml version='1.0' encoding='utf-8'?>
                                <RegistrySet xsi:noNamespaceSchemaLocation='schema.xsd'
                                 xmlns:xsi='http://www.w3.org/2001/XMLSchema-instance'>
                                '''
    protocol = cur.execute(f'SELECT * FROM protocols WHERE id={id}').fetchone()
    program = cur.execute(f'SELECT * FROM studyingPrograms WHERE id={protocol[1]}').fetchone()
    protocolClients = cur.execute(f'SELECT * FROM protocolsClients WHERE idProtocols={id}').fetchall()
    name = f'xml Протокол №{id} от {protocol[3]}'
    
    with open('company.json') as json_file:
        data = json.load(json_file)
        Organization = f'''<Organization>
                        <Inn>{data['inn']}</Inn>
                        <Title>{data['name']}</Title>
                    </Organization>'''
    Test = f'''<Test isPassed="true" learnProgramId="1">
                    <Date>{protocol[3]}</Date>
                    <ProtocolNumber>{id}</ProtocolNumber>
                    <LearnProgramTitle>{protocol[2]}</LearnProgramTitle>
                </Test>'''
    for prot in protocolClients:
        RegistryRecord = "<RegistryRecord>"
        nameCli = prot[2]
        firstSp = nameCli.find(' ')
        LastName = nameCli[:firstSp] if firstSp != -1 else nameCli
        secSp = nameCli[firstSp+1:].find(' ')
        FirstName = nameCli[firstSp+1:secSp] if secSp != -1 else nameCli[firstSp+1:]
        MiddleName = nameCli[secSp+1:] if secSp != -1 else ''
        Snils = prot[4]
        trainee = cur.execute(f'SELECT * FROM trainee WHERE snils={Snils}').fetchone()
        Position = trainee[2] if trainee else ''
        client = cur.execute(f'SELECT * FROM clients WHERE id={prot[3]}').fetchone()
        EmployerInn = client[2]
        EmployerTitle = client[1]
        addWork =   f'''
        <Worker>
            <LastName>{LastName}</LastName>
            <FirstName>{FirstName}</FirstName>
            <MiddleName>{MiddleName}</MiddleName>
            <Snils>{Snils}</Snils>
            <Position>{Position}</Position>
            <EmployerInn>{EmployerInn}</EmployerInn>
            <EmployerTitle>{EmployerTitle}</EmployerTitle>
        </Worker>
                    '''
        RegistryRecord += addWork
        RegistryRecord += Organization
        RegistryRecord += Test
        RegistryRecord += '</RegistryRecord>'
        bodyXml += RegistryRecord
    bodyXml += '</RegistrySet>'
    
    response = HttpResponse(bodyXml, content_type='application/xml')
    response['Content-Disposition'] = 'attachment; filename=' + escape_uri_path(
                f'{name}.xml')
    return response

def protocols(request):
    if not request.COOKIES.get('id'):
        return HttpResponseRedirect('/login')
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    if request.method == 'POST':
        listData = list(request.POST.values())[1].lower()
        return HttpResponseRedirect(f'/protocols/search?request={listData}')
    else:
        form = SearchProtocols(request.POST)
        allProtocols = cur.execute('SELECT * FROM protocols').fetchall()
        nList = []
        for protocol in range(len(allProtocols)):
            listR = list(allProtocols[protocol])
            listR[5] = (cur.execute(f'select name from users where id = {listR[5]}').fetchone()[0])
            nList.append(listR[:-1])
        return render(request, 'protocols.html', {'allProtocols': nList, 'form': form})

def protocolsClients(request):
    z = request.GET.get('id')
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    idProtocols = -1 if z == "new" else int(z)
    def retRender(data):
        form1 = ProtocolsForm(my_arg=data[:3])
        listClient1 = data[3:]
        formsClients = []
        for j in range(len(listClient1)//5):
            formsClients.append(ProtocolsClientsForm(my_arg=listClient1[j * 5:(j + 1) * 5], prefix=j))
        allProtocol = list(cur.execute(f'SELECT * FROM protocols where id = {idProtocols}').fetchone())
        allProtocol[5] = cur.execute(f'select name from users where id = {allProtocol[5]}').fetchone()[0]
        allProtocol = [allProtocol[0]] + allProtocol[3:]
        con.commit()
        return render(request, 'protocolsClients.html',
                      {'form': form1, 'formsClients': formsClients, 'allProtocols': allProtocol})
    def retNewRender(data):
        import datetime
        id = len(cur.execute(f"SELECT * FROM protocols").fetchall())
        if not data:
            form1 = ProtocolsForm()
            listClient1 = data[1:]
            return render(request, 'protocolsClients.html',
                          {'form': form1, 'formsClients': [],
                           'allProtocols': [id, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'), request.COOKIES.get('id'), 0]})
        form1 = ProtocolsForm(my_arg=data[:3])
        listClient1 = data[3:]
        formsClients1 = []
        for j in range(len(listClient1)//4):
            formsClients1.append(ProtocolsClientsForm(my_arg=listClient1[j * 5:(j + 1) * 5], prefix=j))
        con.commit()
        return render(request, 'protocolsClients.html',
                      {'form': form1, 'formsClients': formsClients1,
                       'allProtocols': [id, datetime.datetime.now(settings.TIME_DELTA).strftime('%d.%m.%Y %X'),
                                        request.COOKIES.get('id'), 0]})
    def addLine(listData1):
        clientID = \
        cur.execute(f"SELECT idClient FROM applications WHERE idTrainee = {listTrainee[trainee][0]}").fetchone()[0]
        for el in [listTrainee[trainee][1], clientID, listTrainee[trainee][4], None]:
            listData1.append(el)
        return listData1
    if request.method == 'POST':
        listData = list(request.POST.values())[1:]
        value = str(listData[0])
        listData = listData[1:]
        protocol = [listData[0]]
        if value == "0":
            nameClient = cur.execute(f'SELECT name, textForReport FROM studyingPrograms where id = {protocol[0]}').fetchone()
            protocol.append(nameClient[0])
            des = listData[1]
            if listData[1] == '':
                des = nameClient[1]
            protocol.append(des)
            if idProtocols == -1:
                idProtocols = len(cur.execute(f"SELECT * FROM protocols").fetchall())
                listData1 = tuple(
                    protocolsTranslate(([idProtocols] + listData[:1]), request.COOKIES.get('id'))+
                    [listData[1]]
                )
                cur.execute(f'''INSERT INTO protocols VALUES(?,?,?,?,?,?,?)''', listData1)
                countTrainee = 0
                listData = listData[3:]
                for i in range(len(listData) // 5):
                    listValue = listData[i * 5:(i + 1) * 5]
                    idr = \
                    cur.execute(f'SELECT MAX(id) FROM protocolsClients WHERE idProtocols = {idProtocols}').fetchone()[0]
                    idTrainee = int(idr if isinstance(idr, int) else -1) + 1
                    
                    cur.execute(f'INSERT INTO protocolsClients values' +
                                f"({idTrainee}, {idProtocols}, '{listValue[1]}', {listValue[2]}, '{listValue[3]}',"
                                f" '{listValue[4]}', NULL, '{listValue[0]}')")
                    countTrainee += 1
                cur.execute(f"update protocols set countTrainee = {countTrainee} WHERE id = {idProtocols}")
            else:
                cur.execute(
                    f'update protocols set idStudyingProgram = ?, nameStudyingProgram = ?,'
                    f' description = ?'
                    f' where id = {idProtocols}',
                    tuple(protocol)
                )
                
                cur.execute(f"DELETE FROM protocolsClients WHERE idProtocols={idProtocols}")
                countTrainee = 0
                listData = listData[3:]
                for i in range(len(listData) // 5):
                    listValue = listData[i * 5:(i + 1) * 5]
                    idr = cur.execute(f'SELECT MAX(id) FROM protocolsClients WHERE idProtocols = {idProtocols}').fetchone()[0]
                    idTrainee = int(idr if isinstance(idr, int) else -1) + 1
                    cur.execute(f"INSERT INTO protocolsClients values({idTrainee}, {idProtocols},"
                                f" '{listValue[1]}', '{listValue[2]}', '{listValue[3]}', '{listValue[4]}',"
                                f" NULL, '{listValue[0]}')")
                    countTrainee += 1
                cur.execute(f"update protocols set countTrainee = {countTrainee} WHERE id = {idProtocols}")
            con.commit()
            return HttpResponseRedirect(f'/protocols/clients?id={idProtocols}')
        elif value == "1":
            for _ in range(5):
                listData.append(None)
        elif value == "2":
            idStudyingProgram = protocol[0]
            listTrainee = cur.execute(f"SELECT * FROM trainee WHERE prefStudyingProgram = {idStudyingProgram}").fetchall()
            listData1 = []
            for trainee in range(len(listTrainee)):
                tf = True
                traineeProtocols = cur.execute(f"SELECT * FROM protocolsClients WHERE fullName = "
                                                +f"'{listTrainee[trainee][1]}' and snils = '{listTrainee[trainee][4]}' ").fetchall()
                if not traineeProtocols:
                    listData1 = addLine(listData1)
                else:
                    for traineeOne in traineeProtocols:
                        id2 = cur.execute(f"SELECT idStudyingProgram FROM protocols WHERE id = {traineeOne[1]}").fetchone()[0]
                        if str(id2) == str(idStudyingProgram):
                            tf = False
                    if tf:
                        listData1 = addLine(listData1)
            if listData1:
                listData = listData[:3]
                for el in listData1:
                    listData.append(el)
        elif value == "3":
            listData1 = listData[3:]
            for i in range(len(listData1) // 5):
                listValue2 = listData1[i * 5:(i + 1) * 5]
                listValue2[0] = f"{listData[2]}{i+1:02d}"
                listData1[i * 5:(i + 1) * 5] = listValue2
            listData = listData[:3] + listData1
        else:
            idDel = int(value)-4
            listData = listData[:5*idDel+1] + listData[5*(idDel+1)+1:]
        if idProtocols == -1:
            return retNewRender(listData)
        return retRender(listData)
    else:
        if idProtocols == -1:
            return retNewRender([])
        allProtocols = list(cur.execute(f"SELECT * FROM protocols where id = '{idProtocols}'").fetchone())
        allProtocols[5] = cur.execute(f"select name from users where id = '{allProtocols[5]}'").fetchone()[0]
        allApplicationsForm = [allProtocols[1]]+allProtocols[6:]
        allProtocols = [allProtocols[0]] + allProtocols[3:]
        form = ProtocolsForm(my_arg=allApplicationsForm)
        listClient = list(cur.execute(f'SELECT * FROM protocolsClients where idProtocols = {allProtocols[0]}').fetchall())
        formsClients = []
        for i in range(len(listClient)):
            valueList = [listClient[i][7]] + list(listClient[i][2:7])
            formsClients.append(ProtocolsClientsForm(my_arg=valueList, prefix=i))
        return render(request, 'protocolsClients.html',
                      {'form': form, 'formsClients': formsClients, 'allProtocols': allProtocols})

def company(request):
    import json
    if request.method == 'POST':
        listData = list(request.POST.values())[1:]
        with open('company.json', 'w') as json_file:
            data = {
                "name": listData[0],
                "inn": listData[1],
                "kpp": listData[2],
                "bankAccount": listData[3],
                "address": listData[4],
                "supervisor": listData[5]
                }
            dataJson = json.dumps(data)
            json_file.write(dataJson)
        return HttpResponseRedirect(f'/information/company')
    else:
        with open('company.json') as json_file:
            data = list(json.load(json_file).values())
        form = CompanyForm(my_arg=data)
        return render(request, 'company.html', {'form': form})

def reports(request):
    if request.method == 'POST':
        import datetime, sqlite3, docx
        con = sqlite3.connect('educationalDate.db')
        cur = con.cursor()
        
        listData = list(request.POST.values())[1:]
        if '' in listData:
            form = ReportsForm()
            return render(request, 'reports.html',
                          {'form': form, 'out':'Заполните все поля'})
        dateOt = datetime.datetime.strptime(listData[0], '%Y-%m-%d')
        dateDo = datetime.datetime.strptime(listData[1], '%Y-%m-%d')
        
        protocol = cur.execute(f'SELECT * FROM protocols').fetchall()
        trueProtocols = []
        for one in protocol:
            if dateOt < datetime.datetime.strptime(one[3], '%d.%m.%Y %X') < dateDo:
                groupProt = cur.execute(f'SELECT idGroupReports FROM studyingPrograms WHERE id={one[1]}').fetchone()[0]
                if len(trueProtocols) < groupProt+1:
                    trueProtocols.append([one])
                else:
                    trueProtocols[groupProt].append(one)
        doc = docx.Document()
        h = doc.add_paragraph('Распределение персонала по возрасту и полу')
        h2 = doc.add_paragraph('Распределение персонала без внешних внешних совместителей работающих'
                               ' по договорам гражданско-правового характера по возрасту и полу')
        h.bold = True
        h.alignment = 0
        h2.bold = True
        h2.alignment = 0
        
        #table
        table = doc.add_table(rows=len(trueProtocols) + 3, cols=11)
        table.style = 'Table Grid'
        table.cell(0, 0).merge(table.cell(2, 0))
        table.cell(0, 0).text = ''
        table.cell(0, 1).merge(table.cell(0, 10))
        table.cell(0, 1).text = 'Число полных лет по состоянию на конец отчётного года'
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 1).text = 'моложе 25 лет'
        table.cell(1, 3).merge(table.cell(1, 4))
        table.cell(1, 3).text = '25-29'
        table.cell(1, 5).merge(table.cell(1, 6))
        table.cell(1, 5).text = '30-34'
        table.cell(1, 7).merge(table.cell(1, 8))
        table.cell(1, 7).text = '35-39'
        table.cell(1, 9).merge(table.cell(1, 10))
        table.cell(1, 9).text = '40-44'
        for j in range(1, 10, 2):
            table.cell(2, j).text = 'всего'
            table.cell(2, j+1).text = 'из них женщин'
        for protInd in range(len(trueProtocols)):
            table.cell(3+protInd, 0).text = str(protInd)
            listCount = [0] * 10
            for protInd2 in trueProtocols[protInd]:
                protocolClient2 = cur.execute(f"SELECT * FROM protocolsClients WHERE idProtocols='{protInd2[0]}'").fetchall()
                for protInd1 in protocolClient2:
                    ageGender = cur.execute(f"SELECT age, gender FROM trainee WHERE fullName = "
                                        +f"'{protInd1[2]}' and snils = '{protInd1[4]}' ").fetchone()
                    if not ageGender:
                        continue
                    age = ageGender[0]
                    gender = ageGender[1]
                    
                    if int(age) < 25:
                        listCount[0] += 1
                        if gender == 1:
                            listCount[1] += 1
                    elif int(age) < 30:
                        listCount[2] += 1
                        if gender == 1:
                            listCount[3] += 1
                    elif int(age) < 35:
                        listCount[4] += 1
                        if gender == 1:
                            listCount[5] += 1
                    elif int(age) < 40:
                        listCount[6] += 1
                        if gender == 1:
                            listCount[7] += 1
                    elif int(age) < 45:
                        listCount[8] += 1
                        if gender == 1:
                            listCount[9] += 1
                    for x in range(len(listCount)):
                        table.cell(3 + protInd, x+1).text = str(listCount[x])
        response = HttpResponse(content_type='application/vnd.openxmlformats-'
                                             'officedocument.wordprocessingml.document')
        response['Content-Disposition'] = ("attachment; filename="
                                           + escape_uri_path(
                    f'Отчёт от '
                    f'{dateOt} до {dateDo}.docx'))
        doc.save(response)
        con.commit()
        return response
    else:
        form = ReportsForm()
        return render(request, 'reports.html', {'form': form})

def protocolsSearch(request):
    import sqlite3
    con = sqlite3.connect('educationalDate.db')
    cur = con.cursor()
    if request.method == 'POST':
        listData = list(request.POST.values())[1].lower()
        return HttpResponseRedirect(f'/protocols/search?request={listData}')
    else:
        listData = request.GET.get('request')
        form = SearchProtocols(my_arg=listData)
        protocolsClientsAll = cur.execute('SELECT idProtocols, fullName, idClient FROM protocolsClients').fetchall()
        allListData = []
        for oneHuman in protocolsClientsAll:
            nameClient = cur.execute(f'SELECT name FROM clients WHERE id={oneHuman[2]}').fetchone()
            allListData.append([oneHuman[0]] + list(map(lambda x: x.lower(), list(nameClient))) +
                               list(map(lambda x: x.lower(), oneHuman[1].split())))
        listOut = []
        for allData in range(len(allListData)):
            for w in allListData[allData][1:]:
                if listData == w[:len(listData)]:
                    name = ' '.join(allListData[allData][2:])
                    allListData[allData] = allListData[allData][:2]
                    allListData[allData].append(name)
                    dateProtocol = cur.execute(f"SELECT date FROM protocols WHERE id={allListData[allData][0]}").fetchone()[0]
                    allListData[allData].append(dateProtocol)
                    listOut.append(allListData[allData])
                    break
        listOut = list(listOut)
        return render(request, 'searchProtocols.html',
                      {'form': form, 'listOut': listOut})